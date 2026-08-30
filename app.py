import io
import os
import re
import json
import base64
import fitz  # PyMuPDF – reads text from PDFs
import streamlit as st
import anthropic
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload

# ============================================================
# 1. CONFIG
# ============================================================
MAROON = "#78352A"
LOGO = "ttec_logo.png"
CLAUDE_MODEL = "claude-sonnet-5"
SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]

LANGUAGES = {
    "English": {"code": "English", "rtl": False},
    "Français": {"code": "French", "rtl": False},
    "العربية (Arabic)": {"code": "Arabic", "rtl": True},
}

st.set_page_config(
    page_title="TTEC - Morocco Conformity Skill",
    page_icon=LOGO if os.path.exists(LOGO) else "MA",
    layout="wide",
)

API_KEY = st.secrets["AI_MODEL_API_KEY"]
GOOGLE_DRIVE_FOLDER_ID = st.secrets["DRIVE_FOLDER_ID"]
WORKSPACE_ID = st.secrets.get("ANTHROPIC_WORKSPACE_ID", "")

if WORKSPACE_ID:
    client = anthropic.Anthropic(
        api_key=API_KEY,
        default_headers={"anthropic-workspace-id": WORKSPACE_ID},
    )
else:
    client = anthropic.Anthropic(api_key=API_KEY)


def claude_text(msg):
    """Return the text from a Claude response, skipping any non-text (e.g. thinking) blocks."""
    parts = [b.text for b in msg.content if getattr(b, "type", None) == "text"]
    if not parts:
        parts = [getattr(b, "text", "") for b in msg.content if hasattr(b, "text")]
    return "".join(parts).strip()


# ============================================================
# 1b. OPTIONAL PASSWORD GATE
# ============================================================
# Add  APP_PASSWORD = "your_password"  in Streamlit Secrets to switch this on.
# Leave it out and the app stays open (no gate).
APP_PASSWORD = st.secrets.get("APP_PASSWORD", "")
if APP_PASSWORD:
    if not st.session_state.get("auth_ok"):
        st.title("🔒 TTEC Conformity Skill")
        pw = st.text_input("Enter access password:", type="password")
        if st.button("Enter"):
            if pw == APP_PASSWORD:
                st.session_state["auth_ok"] = True
                st.rerun()
            else:
                st.error("Wrong password.")
        st.stop()

# ============================================================
# 2. STYLING
# ============================================================
st.markdown(f"""
<style>
.stApp {{ background:#ffffff; }}
.ttec-title {{ color:{MAROON}; font-size:1.9rem; font-weight:700; margin:0; }}
.ttec-sub {{ color:#555; margin:.2rem 0 1rem 0; }}
.ttec-rule {{ border:0; border-top:3px solid {MAROON}; margin:.3rem 0 1rem 0; }}
.stButton>button {{ background:{MAROON}; color:#fff; border:none; border-radius:6px; font-weight:600; }}
.stButton>button:hover {{ background:#5c2820; color:#fff; }}
.stDownloadButton>button {{ background:{MAROON}; color:#fff; border:none; border-radius:6px; font-weight:600; }}
div[data-testid="stMarkdownContainer"] table {{ width:100%; border-collapse:collapse; }}
div[data-testid="stMarkdownContainer"] th {{ background:{MAROON}; color:#fff; padding:8px; text-align:left; }}
div[data-testid="stMarkdownContainer"] td {{ border:1px solid #e6dcd7; padding:8px; }}
/* ---- eye-catching report tabs ---- */
button[data-baseweb="tab"] {{ font-weight:700; font-size:1.02rem; color:#9a8f8b; padding:6px 4px; }}
button[data-baseweb="tab"]:hover {{ color:{MAROON}; }}
button[data-baseweb="tab"][aria-selected="true"] {{ color:{MAROON}; }}
[data-baseweb="tab-highlight"] {{ background-color:{MAROON}; height:3px; }}
[data-baseweb="tab-list"] {{ gap:14px; border-bottom:2px solid #eaded9; }}
</style>
""", unsafe_allow_html=True)

# ---------- header ----------
c1, c2 = st.columns([1, 6])
with c1:
    if os.path.exists(LOGO):
        st.image(LOGO, width=170)
with c2:
    st.markdown('<p class="ttec-title">Morocco Product Conformity Skill</p>', unsafe_allow_html=True)
    st.markdown('<p class="ttec-sub">Enter a product or norm codes - applied norm(s), scope, '
                'tests and marking rules, from the TTEC standards library.</p>', unsafe_allow_html=True)
st.markdown('<hr class="ttec-rule">', unsafe_allow_html=True)

# ---------- sidebar ----------
with st.sidebar:
    if os.path.exists(LOGO):
        st.image(LOGO, width=150)
    st.markdown("### How it works")
    st.markdown("1. Choose the **output language**.\n"
                "2. Type a product **or several norm codes**, comma-separated.\n"
                "3. Click **Search standards**.\n"
                "4. Review the relevance tags, **tick** the norms that apply.\n"
                "5. Click **Analyse** - one consolidated report.")
    st.markdown("### Examples")
    st.code("embrayage")
    st.code("joint torique, ISO 3601, ISO 681")
    st.code("jouet, EN 71, EN 62115")
    st.caption("Many products need several norms - add each family as a keyword.")

# ============================================================
# 3. GOOGLE DRIVE
# ============================================================
@st.cache_resource
def get_drive_service():
    creds_dict = dict(st.secrets["gcp_service_account"])
    creds = Credentials.from_service_account_info(creds_dict, scopes=SCOPES)
    return build("drive", "v3", credentials=creds)


def list_all_folder_ids(service, root_id):
    cached = st.session_state.get("folder_ids")
    if cached:
        return cached
    all_ids, to_visit = [root_id], [root_id]
    while to_visit:
        parent = to_visit.pop()
        q = (f"'{parent}' in parents and mimeType = 'application/vnd.google-apps.folder' "
             f"and trashed = false")
        resp = service.files().list(q=q, fields="files(id, name)",
                                    includeItemsFromAllDrives=True,
                                    supportsAllDrives=True).execute()
        for f in resp.get("files", []):
            all_ids.append(f["id"])
            to_visit.append(f["id"])
    st.session_state["folder_ids"] = all_ids
    return all_ids


def _dedupe_key(name):
    """Normalise a filename to the standard's code so duplicates collapse."""
    n = name.lower().rsplit(".pdf", 1)[0]
    n = re.sub(r"\s*\(\d+\)\s*$", "", n)
    n = re.sub(r"[\s_\-]*(copy|copie)\b.*$", "", n)
    n = re.sub(r"[^a-z0-9]", "", n)
    return n or name.lower()


def search_standards(service, keywords):
    folder_ids = list_all_folder_ids(service, GOOGLE_DRIVE_FOLDER_ID)
    matches = []
    for kw in keywords:
        kw = kw.strip()
        if not kw:
            continue
        for i in range(0, len(folder_ids), 15):
            chunk = folder_ids[i:i + 15]
            parents = " or ".join([f"'{fid}' in parents" for fid in chunk])
            q = (f"({parents}) and mimeType = 'application/pdf' and trashed = false "
                 f"and (name contains '{kw}' or fullText contains '{kw}')")
            resp = service.files().list(q=q, fields="files(id, name)",
                                        includeItemsFromAllDrives=True,
                                        supportsAllDrives=True).execute()
            matches.extend(resp.get("files", []))
    seen_id, seen_key, unique = set(), set(), []
    for f in matches:
        if f["id"] in seen_id:
            continue
        seen_id.add(f["id"])
        key = _dedupe_key(f["name"])
        if key in seen_key:
            continue
        seen_key.add(key)
        unique.append(f)
    return unique


def get_text(service, file_id):
    """Download + extract text once, then cache it for the session."""
    cache = st.session_state.setdefault("text_cache", {})
    if file_id in cache:
        return cache[file_id]
    request = service.files().get_media(fileId=file_id)
    buf = io.BytesIO()
    downloader = MediaIoBaseDownload(buf, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    buf.seek(0)
    doc = fitz.open(stream=buf, filetype="pdf")
    text = "".join(page.get_text() for page in doc)
    cache[file_id] = text
    return text


# Boilerplate lines that repeat on every page of IMANOR PDFs (drop them so they don't drown the scope)
_BOILERPLATE = ("accordé sous licence", "licence pour utilisateur", "normalisation@imanor",
                "tous droits réservés", "angle avenue kamal", "droits d'auteur")


def _find_scope_start(low, clean):
    """Find where the real scope section begins, skipping table-of-contents dotted lines."""
    markers = ["domaine d'application", "champ d'application", "objet et domaine",
               "1 objectif", "objectif", "1 objet", "objet", "1 scope", "scope"]
    for m in markers:
        start = 0
        while True:
            idx = low.find(m, start)
            if idx == -1:
                break
            following = clean[idx + len(m): idx + len(m) + 120]
            # a table-of-contents entry looks like "OBJECTIF ........... 3" (many dots)
            if following.count(".") < 8:
                return idx
            start = idx + len(m)
    return -1


def extract_scope(text):
    """Return the Objectif / Domaine d'application section so relevance is judged on real scope."""
    clean = "\n".join(ln for ln in text.splitlines()
                      if not any(b in ln.lower() for b in _BOILERPLATE))
    idx = _find_scope_start(clean.lower(), clean)
    if idx != -1:
        return clean[idx: idx + 2500]
    return clean[:2500]


def read_tech_upload(uploaded):
    """Turn an uploaded technical sheet into usable content.
    Returns (text, image_block). PDF with text -> text. Image or scanned PDF -> image_block
    for Claude vision. Either may be None."""
    if uploaded is None:
        return None, None
    data = uploaded.getvalue()
    name = uploaded.name.lower()
    if name.endswith(".pdf"):
        try:
            doc = fitz.open(stream=data, filetype="pdf")
            text = "".join(p.get_text() for p in doc)
            if len(text.strip()) > 40:            # real text PDF
                return text, None
            # scanned PDF -> render page 1 to PNG for vision
            pix = doc[0].get_pixmap(dpi=150)
            return None, {"media_type": "image/png",
                          "data": base64.b64encode(pix.tobytes("png")).decode()}
        except Exception:
            return None, None
    # image file
    mt = "image/png" if name.endswith(".png") else "image/jpeg"
    return None, {"media_type": mt, "data": base64.b64encode(data).decode()}


# ============================================================
# 4. CLAUDE - RELEVANCE + REPORT
# ============================================================
def suggest_codes(product, language):
    """From a product name/designation, suggest likely codes to search. Keep it short."""
    prompt = f"""Moroccan conformity officer (TTEC). The user has only a product name: "{product}".
Give the standards most likely to apply, SHORT searchable code form (e.g. "EN 71", "EN 62115",
"NM EN 60335", "ISO 3601"). Keep it minimal.
Return ONLY a JSON array (no prose, no fences), max 6 items, each:
{{"code": "<short code>", "reason": "<max 4 words in {language}>"}}"""
    msg = client.messages.create(model=CLAUDE_MODEL, max_tokens=700,
                                 messages=[{"role": "user", "content": prompt}])
    raw = claude_text(msg).replace("```json", "").replace("```", "").strip()
    try:
        data = json.loads(raw)
        return [d for d in data if d.get("code")]
    except Exception:
        return []


def judge_relevance(product, items):
    """items = [(filename, snippet)]. Returns {filename: (verdict, reason)}."""
    listing = "\n\n".join(
        [f"[{i+1}] FILE: {n}\nEXCERPT:\n{s[:1500]}" for i, (n, s) in enumerate(items)]
    )
    prompt = f"""You are a Moroccan conformity officer. The user is looking for the standards
that actually apply to this product: "{product}".

Below are candidate standard documents (filename + opening lines). For EACH document decide how
likely it TRULY governs that product. A document that only mentions the word in passing is NOT
relevant.

{listing}

Return ONLY a JSON array (no prose, no code fences). Each item:
{{"file": "<exact filename>", "verdict": "Likely" or "Maybe" or "Unlikely", "reason": "<one short line>"}}"""
    msg = client.messages.create(model=CLAUDE_MODEL, max_tokens=1500,
                                 messages=[{"role": "user", "content": prompt}])
    raw = claude_text(msg).replace("```json", "").replace("```", "").strip()
    try:
        data = json.loads(raw)
        return {d["file"]: (d.get("verdict", "Maybe"), d.get("reason", "")) for d in data}
    except Exception:
        return {}


def generate_report(product, tech_context, docs, language, tech_image=None):
    per_doc = 40000
    blocks = [f"=== SOURCE DOCUMENT: {name} ===\n{text[:per_doc]}" for name, text in docs]
    combined = "\n\n".join(blocks)
    tech_line = tech_context or ("(see attached technical-sheet image)" if tech_image
                                 else "(none provided)")
    prompt = f"""You are an expert Moroccan market-control and conformity-verification engineer
(TTEC, VOC / PortNet, Law 24-09). You are given one or more official standards that all apply to
the same product. Produce a single CONSOLIDATED verification profile combining them.

TARGET PRODUCT: {product}
USER TECHNICAL DATA: {tech_line}

{combined}

Write the ENTIRE response in {language} (translate the section titles too). Use markdown, with
five sections in this order:

Section 1 - Applied norm(s): SHORT bullet points only, one line per norm (no intro paragraph).
For each, mark GENERAL/BASE, PRODUCT-SPECIFIC, or CONDITIONAL (state the condition). Add referenced
or equivalent standards (EN, ISO, JIS...) as bullets too.

Section 2 - Simplified scope: MAXIMUM 2 short sentences. Then one line exactly like this:
"**In:** <products covered> — **Out:** <products excluded>". Keep it tight.

Section 3 - Mandatory tests: give a SEPARATE table for EACH standard (do not merge them).
Under each standard, put a bold sub-heading with the norm name, then ONE markdown table, 4 columns:
"Applies? | Clause | Test / characteristic | Acceptance criteria".
Fill the "Applies?" column ONLY from the USER TECHNICAL DATA / attached technical sheet:
- 🎯 if the technical data shows this test is relevant to THIS product,
- ➖ if it is clearly not applicable (add 2-3 word reason),
- ❓ if the technical data does not say (verify).
NEVER remove a test row. If no technical data was provided, put "—" in every Applies? cell.

Section 4 - Labelling & marking: give a SEPARATE table for EACH standard (do not merge them).
Under each standard, put a bold sub-heading with the norm name, then ONE markdown table, 3 columns
(required element / placement / language & legibility).

Section 5 - Notes & gaps: AT MOST 6 short bullet points, most important first (no paragraphs).
Flag any missing or garbled value (never invent one); and if a product feature might trigger an
ADDITIONAL norm not among the documents provided (e.g. an electrical toy needing EN 62115), name
it as a bullet reminder.

Use '###' markdown headings for each section title and proper markdown tables.
"""
    content = [{"type": "text", "text": prompt}]
    if tech_image:
        content.append({"type": "image", "source": {
            "type": "base64", "media_type": tech_image["media_type"], "data": tech_image["data"]}})
    msg = client.messages.create(model=CLAUDE_MODEL, max_tokens=6000,
                                 messages=[{"role": "user", "content": content}])
    return claude_text(msg)


def generate_from_knowledge(product, tech_context, language):
    """Fallback: no official document found, answer from general knowledge (UNVERIFIED)."""
    prompt = f"""You are an expert Moroccan market-control and conformity-verification engineer
(TTEC, VOC / PortNet, Law 24-09). NO official standard document was found in the library for this
product, so you must answer FROM GENERAL KNOWLEDGE ONLY - this is an unverified estimate.

TARGET PRODUCT: {product}
USER TECHNICAL DATA: {tech_context or "(none provided)"}

Write the ENTIRE response in {language} (translate the section titles too).
Start with ONE bold warning line, in {language}, stating clearly that this is an UNVERIFIED
general-knowledge estimate, NOT taken from an official document, and that every value must be
checked against the real standard before use.

Then use markdown with these sections:

Section 1 - Likely applicable norm(s): the Moroccan NM and/or EN/ISO standards that most likely
apply (your best estimate). Distinguish base / product-specific / conditional.

Section 2 - Simplified scope: MAXIMUM 2 short sentences, then one line exactly like this:
"**In:** <products usually covered> — **Out:** <products usually excluded>".

Section 3 - Typical mandatory tests: give a SEPARATE table for EACH standard (do not merge).
Under each, a bold norm sub-heading, then ONE table, 4 columns:
"Applies? | Clause | Test / characteristic | Typical criteria".
Fill "Applies?" ONLY from the USER TECHNICAL DATA: 🎯 relevant, ➖ likely not (short reason),
❓ can't tell (verify). Never remove a row. If no technical data was provided, put "—" everywhere.

Section 4 - Typical labelling & marking: give a SEPARATE table for EACH standard. Under each, a
bold norm sub-heading, then ONE table, 3 columns (element / placement / language & legibility).

Section 5 - Cautions: SHORT bullet points only. State what you could NOT confirm, and that the
officer must locate the official NM text plus any decree or email instruction before deciding.

CRITICAL: never invent a precise clause number or numeric threshold. Where you are not sure,
write "verify" instead of a number. Use '###' markdown headings and proper markdown tables.
"""
    msg = client.messages.create(model=CLAUDE_MODEL, max_tokens=6000,
                                 messages=[{"role": "user", "content": prompt}])
    return claude_text(msg)


def generate_coverage(product, standard_docs, language, report_text=None, report_image=None):
    """Compare a lab test report against the mandatory tests of the standard(s)."""
    per_doc = 40000
    blocks = [f"=== STANDARD: {name} ===\n{text[:per_doc]}" for name, text in standard_docs]
    combined = "\n\n".join(blocks)
    rep = (report_text[:30000] if report_text else "(see attached test-report image)")
    prompt = f"""You are a Moroccan conformity officer (TTEC, VOC / PortNet). Compare a laboratory
TEST REPORT against the mandatory tests of the applicable standard(s) for this product: {product}.

APPLICABLE STANDARD(S):
{combined}

TEST REPORT CONTENT:
{rep}

Write the response in {language}, in markdown. Produce ONE table with columns:
"Status | Required test (norm & clause) | Result found in the report | Note".
Status must be exactly one of:
- ✅ COVERED (the test is present in the report),
- ❌ MISSING (required by the standard but not found in the report),
- ⚠️ CHECK (present but the result looks failing, out of tolerance, or unclear).
List EVERY mandatory test of the standard(s) as its own row.
After the table, add one bold summary line: "X covered / Y missing / Z to check".

CRITICAL: you only check what the report DOCUMENT states - you do NOT judge the authenticity or the
technical validity of the laboratory work. End with one bold reminder line, in {language}, that the
report's authenticity must be confirmed directly with the laboratory, and that this is a coverage
check only, not an approval.
"""
    content = [{"type": "text", "text": prompt}]
    if report_image:
        content.append({"type": "image", "source": {
            "type": "base64", "media_type": report_image["media_type"], "data": report_image["data"]}})
    msg = client.messages.create(model=CLAUDE_MODEL, max_tokens=4000,
                                 messages=[{"role": "user", "content": content}])
    return claude_text(msg)


# ============================================================
# 5. WORD EXPORT (with RTL support for Arabic)
# ============================================================
def _set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:bidi"))
    for run in paragraph.runs:
        rPr = run._r.get_or_add_rPr()
        rtl = OxmlElement("w:rtl")
        rtl.set(qn("w:val"), "1")
        rPr.append(rtl)


def _add_runs(paragraph, text, force_bold=False, white=False):
    for part in re.split(r"(\*\*.+?\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        else:
            r = paragraph.add_run(part)
        if force_bold:
            r.bold = True
        if white:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _add_table(doc, table_lines, rtl=False):
    rows = [[c.strip() for c in ln.strip().strip("|").split("|")] for ln in table_lines]
    rows = [r for r in rows if not all(set(c) <= set("-: ") and c != "" for c in r)]
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncol)
    table.style = "Table Grid"
    for ridx, r in enumerate(rows):
        cells = table.add_row().cells
        for cidx in range(ncol):
            txt = r[cidx] if cidx < len(r) else ""
            p = cells[cidx].paragraphs[0]
            if ridx == 0:
                _shade(cells[cidx], "78352A")
                _add_runs(p, txt, force_bold=True, white=True)
            else:
                _add_runs(p, txt)
            if rtl:
                _set_rtl(p)


def report_to_docx(markdown_text, title, rtl=False):
    doc = Document()
    h = doc.add_heading(level=0)
    run = h.add_run(f"TTEC - {title}")
    run.font.color.rgb = RGBColor(0x78, 0x35, 0x2A)
    if rtl:
        _set_rtl(h)

    lines = markdown_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1; continue
        if line.startswith("#"):
            level = min(len(line) - len(line.lstrip("#")), 3)
            hh = doc.add_heading(line.lstrip("#").strip(), level=max(1, level))
            for r in hh.runs:
                r.font.color.rgb = RGBColor(0x78, 0x35, 0x2A)
            if rtl:
                _set_rtl(hh)
            i += 1; continue
        if line.lstrip().startswith("|"):
            tbl = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                tbl.append(lines[i].strip()); i += 1
            _add_table(doc, tbl, rtl=rtl); continue
        if line.lstrip().startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, line.lstrip()[2:])
            if rtl:
                _set_rtl(p)
            i += 1; continue
        p = doc.add_paragraph()
        _add_runs(p, line)
        if rtl:
            _set_rtl(p)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ============================================================
# 6. USER INTERFACE
# ============================================================
lang_label = st.selectbox("Output language / Langue de sortie / لغة النتيجة", list(LANGUAGES.keys()))

# If the code-helper queued a search term, inject it BEFORE the text box is created
if "inject_kw" in st.session_state:
    st.session_state["kw_input"] = st.session_state.pop("inject_kw")

kw_input = st.text_input("Product or norm codes (comma-separated):",
                         placeholder="e.g. jouet, EN 71, EN 62115", key="kw_input")

# --- Code-suggestion helper (when you only have a product name/designation) ---
with st.container():
    if st.button("💡 Suggest norm codes (I only have a product name)"):
        if kw_input.strip():
            with st.spinner("Suggesting likely standard codes..."):
                st.session_state["code_sugg"] = suggest_codes(kw_input.strip(),
                                                              LANGUAGES[lang_label]["code"])
        else:
            st.warning("Type a product name or designation first.")
    sugg = st.session_state.get("code_sugg")
    if sugg:
        st.caption("💡 Likely codes (estimate — verify). These are what to search:")
        for s in sugg:
            st.markdown(f"🔎 **{s.get('code','')}** — {s.get('reason','')}")
        codes_joined = ", ".join(s.get("code", "") for s in sugg if s.get("code"))
        if st.button(f"🔍 Search these codes → {codes_joined[:70]}"):
            st.session_state["inject_kw"] = codes_joined
            st.session_state["do_search"] = True
            st.session_state.pop("code_sugg", None)
            st.rerun()

tech = st.text_area("Technical sheet (optional but recommended) - paste the product spec here to "
                    "flag which tests apply 🎯 / likely not ➖ / verify ❓:")
tech_file = st.file_uploader("...or upload the technical sheet (PDF / JPG / PNG):",
                             type=["pdf", "jpg", "jpeg", "png"])

# --- Step 1: search + relevance ---
if st.button("🔍 Search standards") or st.session_state.pop("do_search", False):
    kws = [k for k in kw_input.split(",") if k.strip()]
    st.session_state.pop("report", None)
    if not kws:
        st.warning("Type at least one product or norm code.")
        st.session_state.pop("files", None)
    else:
        try:
            service = get_drive_service()
            with st.spinner("Searching the TTEC standards library..."):
                files = search_standards(service, kws)
            st.session_state["files"] = files
            st.session_state["product"] = kw_input.strip()
            # merge pasted text + any uploaded technical sheet
            up_text, up_image = read_tech_upload(tech_file)
            merged = "\n".join(t for t in [tech, up_text] if t and t.strip())
            st.session_state["tech"] = merged
            st.session_state["tech_image"] = up_image
            if files:
                st.session_state.pop("trunc_note", None)
            else:
                # No official document in the Drive -> automatic general-knowledge fallback
                with st.spinner("No official document found - preparing a general-knowledge estimate..."):
                    st.session_state["report"] = generate_from_knowledge(
                        kw_input.strip(), tech, LANGUAGES[lang_label]["code"])
                    st.session_state["report_sources"] = []
                    st.session_state["report_rtl"] = LANGUAGES[lang_label]["rtl"]
                    st.session_state["report_unverified"] = True
        except Exception as e:
            st.error(f"Error - check the [gcp_service_account] secret, the folder share, and that "
                     f"the Drive API is enabled. Details: {e}")
            st.session_state.pop("files", None)

# --- Step 2: choose norms + analyse ---
files = st.session_state.get("files")
if files is not None:
    if not files:
        st.info("No official document found in the Drive for this search - showing a "
                "general-knowledge estimate below. Tip: try the code (💡 button) or the NM code.")
    else:
        names = [f["name"] for f in files]
        st.success(f"{len(files)} standard(s) found. Tick the ones to analyse:")
        default = names if len(names) <= 6 else []
        chosen = st.multiselect("Norms to analyse:", names, default=default)
        if st.button("✨ Analyse selected norms") and chosen:
            service = get_drive_service()
            docs, empty = [], []
            with st.spinner("Reading documents..."):
                for f in files:
                    if f["name"] in chosen:
                        t = get_text(service, f["id"])
                        (docs if t.strip() else empty).append((f["name"], t))
            if empty:
                st.warning("No extractable text (likely scans - OCR needed): "
                           + ", ".join(n for n, _ in empty))
            if docs:
                with st.spinner("Generating consolidated report..."):
                    st.session_state["report"] = generate_report(
                        st.session_state["product"], st.session_state.get("tech", ""),
                        docs, LANGUAGES[lang_label]["code"],
                        tech_image=st.session_state.get("tech_image"))
                    st.session_state["report_sources"] = [n for n, _ in docs]
                    st.session_state["analysed_docs"] = docs
                    st.session_state["report_rtl"] = LANGUAGES[lang_label]["rtl"]
                    st.session_state["report_unverified"] = False

# --- Step 3: show report + download ---
report = st.session_state.get("report")
if report:
    st.markdown('<hr class="ttec-rule">', unsafe_allow_html=True)
    unverified = st.session_state.get("report_unverified", False)
    if unverified:
        st.warning("⚠️ UNVERIFIED — general-knowledge estimate, not from an official document. "
                   "No matching standard was found in the Drive. Verify every value against the "
                   "real NM text (and any decree / email instruction) before use.")
    else:
        with st.expander("📄 Documents used in this report"):
            for n in st.session_state.get("report_sources", []):
                st.write("•", n)
    rtl = st.session_state.get("report_rtl")

    # Split the report into its "### " sections and show them as tabs for quick reading
    sections = []
    current_title, current_body = None, []
    for ln in report.splitlines():
        if ln.strip().startswith("###"):
            if current_title is not None:
                sections.append((current_title, "\n".join(current_body)))
            current_title = ln.lstrip("#").strip()
            current_body = []
        else:
            current_body.append(ln)
    if current_title is not None:
        sections.append((current_title, "\n".join(current_body)))

    def short(t):
        t = re.sub(r"^\s*(section\s*)?\d+[\.\)\-]?\s*", "", t, flags=re.I)  # drop "Section 1 -"
        t = t.lstrip(" -–—:•").strip()
        return (t[:20] + "…") if len(t) > 21 else t

    # Fixed order 1..5 -> icons. Scope / Tests / Marking (2,3,4) get vivid icons to catch the eye.
    ICONS = ["📋", "🎯", "🧪", "🏷️", "📝"]

    if len(sections) >= 2:
        labels = [f"{ICONS[i] if i < len(ICONS) else '•'} {short(t)}"
                  for i, (t, _) in enumerate(sections)]
        tabs = st.tabs(labels)
        for tab, (title, body) in zip(tabs, sections):
            with tab:
                block = f"### {title}\n{body}"
                if rtl:
                    st.markdown(f"<div dir='rtl'>{block}</div>", unsafe_allow_html=True)
                else:
                    st.markdown(block)
    else:
        if rtl:
            st.markdown(f"<div dir='rtl'>{report}</div>", unsafe_allow_html=True)
        else:
            st.markdown(report)

    title = st.session_state.get("product", "report")
    fname = "TTEC_conformity_report.docx"
    if unverified:
        title = "[UNVERIFIED estimate] " + title
        fname = "TTEC_UNVERIFIED_estimate.docx"
    st.download_button(
        "⬇️ Download as Word (.docx)",
        report_to_docx(report, title, rtl=st.session_state.get("report_rtl", False)),
        file_name=fname,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

# ============================================================
# 7. TEST REPORT COVERAGE CHECK (optional, self-contained)
# ============================================================
analysed = st.session_state.get("analysed_docs")
if analysed:
    st.markdown('<hr class="ttec-rule">', unsafe_allow_html=True)
    with st.expander("🧪 Test report coverage check (optional)"):
        st.caption("Upload a lab test report to see which required tests are covered ✅ / missing ❌ "
                   "/ to check ⚠️ against the standard(s) above. This reads what the report SAYS - it "
                   "does NOT verify authenticity (confirm that directly with the laboratory).")
        rep_file = st.file_uploader("Test report (PDF / JPG / PNG):",
                                    type=["pdf", "jpg", "jpeg", "png"], key="rep_up")
        rep_paste = st.text_area("...or paste the report text:", key="rep_paste")
        if st.button("Check coverage"):
            rtxt, rimg = read_tech_upload(rep_file)
            merged = "\n".join(t for t in [rep_paste, rtxt] if t and t.strip())
            if not merged and not rimg:
                st.warning("Upload or paste a test report first.")
            else:
                with st.spinner("Comparing the report against the standard(s)..."):
                    st.session_state["coverage"] = generate_coverage(
                        st.session_state.get("product", ""), analysed,
                        LANGUAGES[lang_label]["code"], report_text=merged or None,
                        report_image=rimg)

        cov = st.session_state.get("coverage")
        if cov:
            if st.session_state.get("report_rtl"):
                st.markdown(f"<div dir='rtl'>{cov}</div>", unsafe_allow_html=True)
            else:
                st.markdown(cov)
            st.download_button(
                "⬇️ Download coverage (.docx)",
                report_to_docx(cov, "Coverage - " + st.session_state.get("product", "report"),
                               rtl=st.session_state.get("report_rtl", False)),
                file_name="TTEC_coverage_check.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="cov_dl")
